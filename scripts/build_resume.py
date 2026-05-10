#!/usr/bin/env python3
"""
build_resume.py
Converts a resume markdown file to a .docx by using the reference resume as a
style template, clearing its body, and rebuilding with parsed content.

This approach preserves every style, numbering, theme font, and border definition
from the reference without any guesswork.

Usage: uv run python scripts/build_resume.py <resume.md> [output.docx]
"""

import sys
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

REPO_ROOT      = Path(__file__).resolve().parent.parent
TEMPLATE_DOCX  = REPO_ROOT / "resume" / "Gaurav_Resume_1_2025-11-16.docx"

# US Letter page size
PAGE_WIDTH  = Inches(8.5)
PAGE_HEIGHT = Inches(11.0)
MARGIN      = Inches(0.5)

# Right tab stop position for dates (page width minus both margins, in twips)
# 8.5" - 0.5" - 0.5" = 7.5" = 10800 twips
RIGHT_TAB_TWIPS = int(7.5 * 1440)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def load_template(template_path: Path) -> Document:
    """Open the reference docx and clear its body, keeping all style/numbering/theme data."""
    doc = Document(template_path)
    body = doc.element.body

    # Keep the section properties element (page size, margins, etc.)
    sect_pr = body.find(qn("w:sectPr"))

    # Remove all body children
    for child in list(body):
        body.remove(child)

    # Re-add sectPr as the only body element
    if sect_pr is not None:
        body.append(sect_pr)

    # Update page size to US Letter (template is A4)
    section = doc.sections[0]
    section.page_width  = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin    = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin   = MARGIN
    section.right_margin  = MARGIN

    return doc


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def set_right_tab(paragraph, pos_twips: int = RIGHT_TAB_TWIPS) -> None:
    """Add a right-aligned tab stop at pos_twips to the paragraph."""
    pPr  = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def set_space_before(paragraph, pt: float) -> None:
    """Set space-before on a paragraph in points."""
    pPr  = paragraph._p.get_or_add_pPr()
    pSpc = paragraph._p.pPr.get_or_add_pPr() if hasattr(paragraph._p, "pPr") else OxmlElement("w:spacing")
    pSpc = OxmlElement("w:spacing")
    pSpc.set(qn("w:before"), str(int(pt * 20)))
    pSpc.set(qn("w:beforeAutospacing"), "0")
    pPr.append(pSpc)


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

def strip_links(text: str) -> str:
    """[label](url) → label; un-escape common markdown escapes."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace(r"\+", "+").replace(r"\&", "&").replace(r"\-", "-")


def add_inline_runs(paragraph, text: str,
                    base_bold: bool = False, base_italic: bool = False) -> None:
    """
    Parse **bold** and *italic* inline markdown and append correctly styled
    runs to the paragraph. Font properties inherit from the paragraph style.
    """
    text = strip_links(text)
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)", text, re.DOTALL):
        bold_t, italic_t, plain_t = m.group(1), m.group(2), m.group(3)
        if bold_t is not None:
            run = paragraph.add_run(bold_t)
            run.bold   = True
            run.italic = base_italic
        elif italic_t is not None:
            run = paragraph.add_run(italic_t)
            run.bold   = base_bold
            run.italic = True
        elif plain_t:
            run = paragraph.add_run(plain_t)
            run.bold   = base_bold
            run.italic = base_italic


# ---------------------------------------------------------------------------
# Paragraph builders  (each mirrors a style from the reference docx)
# ---------------------------------------------------------------------------

def add_name(doc: Document, text: str) -> None:
    """'Title' style — 16pt Calibri Light bold, centered."""
    p = doc.add_paragraph(style="Title")
    add_inline_runs(p, text, base_bold=True)


def add_contact(doc: Document, text: str) -> None:
    """'Normal' style, explicitly centered — contact info line."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, strip_links(text))


def add_section_heading(doc: Document, text: str) -> None:
    """'Heading 1' style — bold, left-aligned, 6pt space before, with bottom border."""
    p = doc.add_paragraph(style="Heading 1")
    add_inline_runs(p, text, base_bold=True)


def add_role_heading(doc: Document, text: str) -> None:
    """'Heading 2' style — bold, left/right tab split for title and date."""
    p = doc.add_paragraph(style="Heading 2")
    set_right_tab(p)
    parts = text.split("\t", 1)
    if len(parts) == 2:
        add_inline_runs(p, parts[0], base_bold=True)
        p.add_run("\t")
        add_inline_runs(p, parts[1], base_bold=True)
    else:
        add_inline_runs(p, text, base_bold=True)


def add_company(doc: Document, text: str) -> None:
    """'Heading 3' style — italic."""
    p = doc.add_paragraph(style="Heading 3")
    add_inline_runs(p, text, base_italic=True)


def add_bullet(doc: Document, text: str) -> None:
    """'List Paragraph' style — inherits bullet from style's numPr definition."""
    p = doc.add_paragraph(style="List Paragraph")
    add_inline_runs(p, text)


def add_education_line(doc: Document, text: str) -> None:
    """'Normal' style, right-aligned tab for the date column."""
    p = doc.add_paragraph(style="Normal")
    set_right_tab(p)
    if "\t" in text:
        left, right = text.split("\t", 1)
        add_inline_runs(p, left)
        p.add_run("\t")
        add_inline_runs(p, right)
    else:
        add_inline_runs(p, text)


def add_plain_line(doc: Document, text: str) -> None:
    """'Normal' style for any other plain or inline-formatted line."""
    p = doc.add_paragraph(style="Normal")
    add_inline_runs(p, text)


# ---------------------------------------------------------------------------
# Line classifiers
# ---------------------------------------------------------------------------

_SECTION_RE = re.compile(
    r"^# \*\*(SUMMARY|PROFESSIONAL|EDUCATION|FSAE|PROJECTS|SOFTWARE)", re.I
)

def is_name_line(s: str) -> bool:
    return bool(re.match(r"^# \*\*.+\*\*$", s)) and not _SECTION_RE.match(s)

def is_section_heading(s: str) -> bool:
    return bool(_SECTION_RE.match(s))

def is_role_heading(s: str) -> bool:
    return s.startswith("## **")

def is_company_heading(s: str) -> bool:
    return bool(re.match(r"^### \*.+\*$", s))

def is_bullet(s: str) -> bool:
    return s.startswith("* ")

def is_contact_line(s: str) -> bool:
    return ("|" in s
            and ("@" in s or "linkedin" in s.lower() or "github" in s.lower())
            and not s.startswith("#")
            and not s.startswith("*"))


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def build_docx(md_path: Path, docx_path: Path,
               template_path: Path = TEMPLATE_DOCX) -> None:

    doc = load_template(template_path)

    for line in md_path.read_text(encoding="utf-8").splitlines():
        s = line.strip().rstrip("\\")  # strip pandoc hard-break marker
        if not s:
            continue

        if is_name_line(s):
            name = re.sub(r"^# \*\*(.+)\*\*$", r"\1", s)
            add_name(doc, name)

        elif is_contact_line(s):
            add_contact(doc, s)

        elif is_section_heading(s):
            heading = re.sub(r"^# \*\*(.+)\*\*$", r"\1", s)
            add_section_heading(doc, heading)

        elif is_role_heading(s):
            role = re.sub(r"^## \*\*(.+)\*\*$", r"\1", s)
            add_role_heading(doc, role)

        elif is_company_heading(s):
            company = re.sub(r"^### \*(.+)\*$", r"\1", s)
            add_company(doc, company)

        elif is_bullet(s):
            add_bullet(doc, s[2:])

        elif s.startswith("**") or s.startswith("*") or s.startswith("["):
            # Education entries and similar inline-formatted lines
            add_education_line(doc, s)

        else:
            add_plain_line(doc, s)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: build_resume.py <resume.md> [output.docx]", file=sys.stderr)
        sys.exit(1)

    md  = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    build_docx(md, out)
